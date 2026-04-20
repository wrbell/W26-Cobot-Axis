#!/usr/bin/env python3
"""Render reports/turn-in/report/report.md -> report.docx via python-docx.

Keeps Markdown as the source of truth. Regenerate whenever report.md changes.
Prints per-section body word counts (Sections A-J only, excluding headings,
tables, figure captions, HTML comments, and references) to help stay under
the 2000-word limit.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = Path(__file__).parent
SRC = HERE / "report.md"
OUT = HERE / "report.docx"

HTML_COMMENT = re.compile(r"<!--.*?-->", re.DOTALL)
INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*([^*]+?)\*(?!\*)")
INLINE_CODE = re.compile(r"`([^`]+?)`")
INLINE_LINK = re.compile(r"\[([^\]]+?)\]\(([^)]+?)\)")


def strip_comments(text: str) -> str:
    return HTML_COMMENT.sub("", text)


def add_runs(paragraph, text: str) -> None:
    """Render Markdown inline formatting into python-docx runs."""
    # Links -> just use the link text (plain); URLs belong in a references section
    text = INLINE_LINK.sub(r"\1", text)

    # Walk bold/italic/code tokens in order
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        else:  # *italic*
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    """Parse a GitHub-style pipe table into rows (dropping the separator row)."""
    rows: list[list[str]] = []
    for line in lines:
        if set(line.replace("|", "").replace(":", "").replace("-", "").strip()) == set():
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def is_bullet(line: str) -> bool:
    return line.lstrip().startswith(("- ", "* ", "+ "))


def word_count(text: str) -> int:
    return len([w for w in re.split(r"\s+", text.strip()) if w])


def main() -> int:
    if not SRC.exists():
        print(f"ERROR: {SRC} not found", file=sys.stderr)
        return 1

    raw = SRC.read_text(encoding="utf-8")
    text = strip_comments(raw)
    lines = text.splitlines()

    doc = Document()

    # Default body font a bit tighter for academic prose
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Accumulators for per-section word counts (body paragraphs only, excluding
    # tables, headings, and the appendix/references). We count each section
    # until the next H1/H2 section marker.
    section_word_counts: dict[str, int] = {}
    current_section = "(preamble)"
    # Flag: don't count words once we enter References or Appendix
    count_words = True

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # Table block
        if is_table_line(line):
            block: list[str] = []
            while i < n and is_table_line(lines[i]):
                block.append(lines[i])
                i += 1
            rows = parse_table_rows(block)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r_idx, row_cells in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_cells):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_runs(p, cell_text)
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
            doc.add_paragraph()  # spacing after table
            continue

        # Headings
        if line.startswith("# "):
            heading_text = line[2:].strip()
            current_section = heading_text
            count_words = True
            doc.add_heading(heading_text, level=1)
            i += 1
            continue
        if line.startswith("## "):
            heading_text = line[3:].strip()
            current_section = heading_text
            # References and Appendix are excluded from the body word count
            if heading_text.lower().startswith(("references", "appendix")):
                count_words = False
            else:
                count_words = True
            doc.add_heading(heading_text, level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Horizontal rule -> skip (we use headings for structure)
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # Bullet list
        if is_bullet(line):
            bullet_lines: list[str] = []
            while i < n and is_bullet(lines[i]):
                bullet_lines.append(lines[i].lstrip()[2:])
                i += 1
            for b in bullet_lines:
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, b)
                if count_words:
                    section_word_counts[current_section] = (
                        section_word_counts.get(current_section, 0) + word_count(b)
                    )
            continue

        # Blank line -> skip (python-docx handles spacing via heading styles)
        if not line.strip():
            i += 1
            continue

        # Regular paragraph (may span multiple lines until blank)
        para_lines: list[str] = []
        while i < n and lines[i].strip() and not is_table_line(lines[i]) and not is_bullet(lines[i]) \
                and not lines[i].startswith(("# ", "## ", "### ")) \
                and lines[i].strip() not in ("---", "***", "___"):
            para_lines.append(lines[i])
            i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        if para_text:
            p = doc.add_paragraph()
            add_runs(p, para_text)
            if count_words:
                section_word_counts[current_section] = (
                    section_word_counts.get(current_section, 0) + word_count(para_text)
                )

    doc.save(OUT)

    # Report word counts
    total = sum(section_word_counts.values())
    print(f"\nWrote {OUT.name} ({OUT.stat().st_size} bytes)")
    print("\nBody word count by section (excludes headings, tables, References, Appendix):")
    for section, count in section_word_counts.items():
        bar = "#" * (count // 20)
        print(f"  {count:>4}  {section[:50]:<50} {bar}")
    limit = 2000
    status = "OK" if total <= limit else "OVER"
    print(f"\n  Total body words: {total} / {limit}  [{status}]")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
